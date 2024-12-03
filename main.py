import matplotlib
from flask import Flask, request, render_template, session, url_for, redirect, send_from_directory, send_file, flash
import openpyxl
import psycopg2
import file1
from flask_bcrypt import Bcrypt
import random
import json
import math
from flask_mail import Mail, Message
from datetime import datetime
from file1 import generate_token, get_email_by_token
from docx import Document

from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
import smtplib
import matplotlib.pyplot as plt
matplotlib.use('Agg')
from io import BytesIO
import base64
from flask import render_template


app = Flask(__name__)
bcrypt = Bcrypt(app)
app.secret_key = '1234567890'

MAIL_SERVER = 'smtp.gmail.com'
MAIL_PORT = 587
MAIL_USE_TLS = True
MAIL_USE_SSL = False
MAIL_USERNAME = 'annabarsh12@gmail.com'
MAIL_PASSWORD = 'kacb yjws npcr pcxf'
MAIL_DEFAULT_SENDER = 'your_email@example.com'

mail = Mail(app)
# from filters import setup_environment
# # настройка фильтров
# setup_environment(app)

@app.route('/authentication')
def authentication():
    return render_template('authentication.html')


@app.route('/')
def home():
    return render_template('home.html')


@app.route('/registration')
def registration():
    return render_template('registration.html')


@app.route('/main')
def main():
    return render_template('main.html')


######


######

# @app.route('/take_test')
# def take_test():
#     return render_template('take_test.html')
# @app.route('/take_test')
# def take_test():
#     questions = file1.get_tests_from_db()  # Используем ваш метод для получения вопросов
#     return render_template('take_test.html', questions=questions)






def send_email(sender, recipient, subject, message):
    smtp_server = 'smtp.gmail.com'
    smtp_port = 587
    smtp_username = 'annabarsh12@gmail.com'
    smtp_password = 'kacb yjws npcr pcxf'
    # Создание объекта MIMEMultipart для создания письма
    msg = MIMEMultipart()
    msg['From'] = sender
    msg['To'] = recipient
    msg['Subject'] = subject
    # Добавление текста сообщения в письмо
    msg.attach(MIMEText(message, 'plain'))
    with smtplib.SMTP(smtp_server, smtp_port) as server:
        server.starttls()
        server.login(smtp_username, smtp_password)
        server.send_message(msg)
    return





# @app.route('/forget_password')
# def forget_password():
#     return render_template('forget_password.html')
#
#
# @app.route('/reset_password', methods=['GET', 'POST'])
# def reset_password():
#     if request.method == 'POST':
#         email = request.form['email']
#
#         # Проверьте, существует ли электронная почта в базе данных
#         if file1.check_user_email_in_db(email):
#             # Если электронная почта существует, генерируйте уникальный токен
#             reset_token = generate_token(email)
#             # Стройте ссылку с токеном
#             reset_link = url_for('change_password', token=reset_token, _external=True)
#
#             # reset_link = url_for('new_password', token=reset_token, _external=True)
#
#             # Отправьте электронное письмо на введенный адрес
#             send_reset_email(email, reset_link)
#
#             # Сообщите пользователю, что письмо с инструкциями отправлено
#             flash('Письмо с инструкциями по восстановлению пароля отправлено на вашу почту.', 'info')
#
#             # Перенаправьте пользователя на страницу входа или другую подходящую страницу
#             return redirect(url_for('email_change_password'))
#         else:
#             # Если электронная почта не существует в базе данных, сообщите пользователю об ошибке
#             flash('Электронная почта не найдена в базе данных.', 'error')
#
#     return render_template('forget_password.html')
#
#
# def send_reset_email(email, reset_link):
#     subject = 'Восстановление пароля'
#     recipients = [email]
#     body = f'Для восстановления пароля перейдите по ссылке: {reset_link}'
#
#     # Устанавливаем кодировку явно (UTF-8)
#     message = Message(subject=subject, recipients=recipients, body=body, charset='utf-8')
#
#     try:
#         mail.send(message)
#         print(f'Email sent successfully to {email}.')
#     except Exception as e:
#         print(f'Error sending email to {email}: {e}')
#
#
#
# @app.route('/email_change_password', methods=['GET', 'POST'])
# def email_change_password():
#     if request.method == 'POST':
#         # Получите токен из параметров запроса
#         reset_token = request.args.get('token')
#
#         # Получите адрес электронной почты по токену
#         email = get_email_by_token(reset_token)
#
#         if email:
#             # Если токен действителен, обработайте данные формы
#             new_password = request.form['password']
#
#             # Обновите пароль пользователя в базе данных (замените эту строку своей логикой)
#             update_user_password(email, new_password)
#
#             # Сообщите пользователю, что пароль обновлен успешно
#             flash('Ваш пароль успешно изменен. Вы можете войти, используя новый пароль.', 'success')
#
#             # Перенаправьте пользователя на страницу входа или другую подходящую страницу
#             return redirect(url_for('login'))
#         else:
#             # Если токен недействителен или не найден, выведите сообщение об ошибке
#             flash('Недействительная или устаревшая ссылка для сброса пароля.', 'error')
#             return redirect(url_for('login'))
#
#     # Отобразите страницу ввода нового пароля
#     return render_template('new_password.html')
#
#
# def update_user_password(email, new_password):
#     conn = psycopg2.connect(user="postgres", password="rootparol", host="127.0.0.1", port="5432", database="MyDB")
#     cursor = conn.cursor()
#
#     # Хешируем новый пароль перед сохранением в базу данных
#     hashed_password = bcrypt.generate_password_hash(new_password).decode('utf-8')
#
#     # Обновляем пароль в базе данных
#     query = "UPDATE users SET userpassword = %s WHERE email = %s"
#     cursor.execute(query, (hashed_password, email))
#
#     conn.commit()
#     cursor.close()
#     conn.close()






@app.route('/login', methods=['POST'])
def login():
    username = request.form['username']
    password = request.form['password']
    # username = session.get('username')
    current_stud_id = file1.get_studid_by_username(username)

    session['current_stud_id'] = current_stud_id  # Добавляем в сессию

    if username == 'Admin' and password == '87654321':
        # Возвращаем редирект на страницу admin.html
        return render_template('admin.html')

    conn = psycopg2.connect(user="postgres", password="rootparol", host="127.0.0.1", port="5432",
                            database="MyDB")
    cursor = conn.cursor()
    cursor.execute(f"""SELECT userpassword, username FROM users WHERE username = '{username}' """)
    user_in_db = cursor.fetchone()
    cursor.close()
    conn.close()
    if user_in_db and bcrypt.check_password_hash(user_in_db[0], password):
        ###
        # Вставка новой строки в таблицу

        sender = 'annabarsh12@gmail.com'  # Ваш адрес электронной почты
        recipient = 'annabarsh12@gmail.com'  # Адрес получателя
        subject = 'регистрация'
        message = ('Поздравляю! Вы успешно вошли в систему. Теперь у Вас есть доступ ко всем возможностям и ресурсам, предоставляемым нашей Системой студентов. ')
        send_email(sender, recipient, subject, message)
        ###
        session['username'] = username
        return render_template('main.html')
    else:
        error = []
        if not user_in_db:
            error.append('username')
        if user_in_db and not bcrypt.check_password_hash(user_in_db[0], password):
            error.append('password')
        return render_template('authentication.html', errors=error, username=username, password=password)

    # if len(password) < 8:
    #     return render_template('authentication.html', warning='Пароль должен содержать минимум 8 символов.')

    print(f"""Вошли с данными {username},{password}""")

    # return f'Вход завершен: Имя пользователя: {username}, Пароль: {password}'


@app.route('/register', methods=['POST'])
def register():
    username = request.form['username']
    email = request.form['email']
    password = request.form['password']
    firstname = request.form['firstname']  # Новые поля
    lastname = request.form['lastname']  # Новые поля
    course = request.form['course']  # Новые поля
    groupnumber = request.form['groupnumber']  # Новые поля

    # Проверка информации пользователя перед добавлением в базу данных
    error = file1.check_user_info_before_db(username, password)
    if len(error) > 0:
        return render_template('registration.html', email=email, username=username, password=password, errors=error)
    else:
        # Хеширование пароля перед сохранением в базе данных
        hashed_password = bcrypt.generate_password_hash(password).decode('utf-8')

        # Использование параметризованного запроса для предотвращения SQL-инъекции
        conn = psycopg2.connect(user="postgres", password="rootparol", host="127.0.0.1", port="5432", database="MyDB")
        cursor = conn.cursor()

        query_users = "INSERT INTO users (username, email, userpassword) " \
                      "VALUES (%s, %s, %s) RETURNING userid"

        cursor.execute(query_users, (username, email, hashed_password))
        userid = cursor.fetchone()[0]

        # Вставка данных в таблицу students
        query_students = "INSERT INTO students (firstname, lastname, course, groupnumber, username) " \
                         "VALUES (%s, %s, %s, %s, %s) RETURNING studid"

        cursor.execute(query_students, (firstname, lastname, course, groupnumber, username))
        studid = cursor.fetchone()[0]

        # Вставка данных в таблицу users_stud
        query_insert_users_stud = "INSERT INTO users_stud (userid, studid, username) VALUES (%s, %s, %s)"
        cursor.execute(query_insert_users_stud, (userid, studid, username))

        conn.commit()
        cursor.close()
        conn.close()
        session['username'] = username
        return render_template('main.html')


# @app.route('/profile')
# def profile():
#     # Получаем логин из сессии
#     username = session.get('username')
#
#     # Получаем информацию о студенте из базы данных
#     student_info = file1.get_student_info(username)
#
#     # Передаем информацию в шаблон
#     student_info_dict = {}
#     for key, value in student_info.items():
#         student_info_dict[key] = value
#
#     # Передать словарь в шаблон
#     return render_template('profile.html', student_info=student_info_dict)
#
#     # return render_template('profile.html', text='Дополнительный текст, если необходимо')


@app.route('/profile')
def profile():
    # Проверяем, что пользователь аутентифицирован
    # if 'username' not in session:
    #     # Если не аутентифицирован, перенаправляем на страницу входа
    #     return redirect(url_for('authentication'))

    # Получаем логин из сессии
    username = session.get('username')
    print(username)

    # Получаем информацию о студенте из базы данных
    student_info = file1.get_student_info(username)
    print(student_info)

    # Передаем информацию в шаблон
    student_info_dict = {}
    for key, value in student_info.items():
        student_info_dict[key] = value
    print(student_info_dict)
    # Передать словарь в шаблон
    return render_template('profile.html', student_info=student_info_dict)


@app.route('/professor')
def professor():
    # Проверяем, авторизован ли профессор
    if 'username' in session and session['role'] == 'professor':
        # Если профессор авторизован, получаем информацию о нем и передаем в шаблон
        professor_info = file1.get_professor_info(session['username'])
        return render_template('professor.html', professor_info=professor_info)
    else:
        # Если профессор не авторизован, перенаправляем на страницу входа
        return redirect(url_for('authentication'))


@app.route("/change_username", methods=['POST'])
def change_username():
    new_username = request.form['new-username']
    if file1.check_username(new_username):
        print(session.get('username'))
        file1.execute_query(f"""update users
              set username='{new_username}'
        where username ='{session.get('username')}'""")
        print(session.get('username'))
        file1.execute_query(f"""update students
                     set username='{new_username}'
               where username ='{session.get('username')}'""")
        session['username'] = new_username
        return render_template('profile.html', success='username',
                               student_info=file1.get_data_for_profile_redirect(session.get('username')))

    else:
        errors = []
        if not file1.check_username(new_username):
            errors.append('username')
            print("errors")
            print(errors)
        return render_template('profile.html',
                               errors=errors, show_username_form=True, new_username=new_username,
                               student_info=file1.get_data_for_profile_redirect(session.get('username')))


@app.route('/change_email', methods=['POST'])
def change_email():
    email = request.form['new-email']
    print(email)
    if file1.check_user_email_in_db(email) == 0:
        file1.execute_query(f"""update users
              set email='{email}'
        where username='{session.get('username')}'""")
        return render_template('profile.html', success='email',
                               student_info=file1.get_data_for_profile_redirect(
                                   session.get('username')))  # Перенаправляем пользователя на страницу профиля
    else:
        errors = []
        if file1.check_user_email_in_db(email):
            errors.append('email')

        return render_template('profile.html',
                               errors=errors, show_email_form=True, email=email,
                               student_info=file1.get_data_for_profile_redirect(session.get('username')))


@app.route('/change_password', methods=['POST'])
def change_password():
    old_password = request.form['old-password']
    new_password = request.form['new-password']
    confirm_password = request.form['confirm-password']

    # Получение хешированного пароля из базы данных
    hashed_password_from_db = file1.get_user_password_from_db(session.get('username'))

    if bcrypt.check_password_hash(hashed_password_from_db, old_password) and file1.check_password(
            new_password) and file1.check_confirm_password(new_password, confirm_password):
        # Хеширование нового пароля
        hashed_new_password = bcrypt.generate_password_hash(new_password).decode('utf-8')

        file1.execute_query(f"""UPDATE users
      SET userpassword='{hashed_new_password}'
      WHERE username='{session.get('username')}'""")

        return render_template('profile.html', success='password',
                               student_info=file1.get_data_for_profile_redirect(
                                   session.get('username')))
    else:
        errors = []
        if not bcrypt.check_password_hash(hashed_password_from_db, old_password):
            errors.append('old_password')
        if not file1.check_password(new_password):
            errors.append('new_password')
        if not file1.check_confirm_password(new_password, confirm_password):
            errors.append('confirm_password')

        return render_template('profile.html', errors=errors, show_password_form=True,
                               old_password=old_password,
                               new_password=new_password,
                               confirm_password=confirm_password,
                               student_info=file1.get_data_for_profile_redirect(session.get('username'))
                               )


# Дополнение маршрута /take_test
@app.route('/take_test', methods=['GET', 'POST'])
def take_test():
    # Получаем логин из сессии
    username = session.get('username')
    current_stud_id = file1.get_studid_by_username(username)

    session['current_stud_id'] = current_stud_id  # Добавляем в сессию

    if request.method == 'POST':
        selected_theme = request.form.get('test_theme')
        current_test_id = file1.get_tetstid_by_theme(selected_theme)
        session['current_test_id'] = current_test_id  # Добавляем в сессию

    if request.method == 'POST':
        selected_theme = request.form.get('test_theme')
        current_test_id = file1.get_tetstid_by_theme(selected_theme)
        session['theme'] = selected_theme

        return redirect(url_for('test_form', theme=selected_theme))
    else:
        themes = file1.get_themes_from_db()
        return render_template('take_test.html', themes=themes)


# from jinja2 import Environment
#
# def zip_lists(a, b):
#     return zip(a, b)
#
# env = Environment()
# env.filters['zip_lists'] = zip_lists


@app.route('/test_form')
def test_form():
    selected_theme = request.args.get('theme')

    # Получение вопросов, их ID и ответов к вопросам
    questions_with_ids = file1.get_questions_from_db(selected_theme)
    questions = [q[1] for q in questions_with_ids]
    question_ids = [q[0] for q in questions_with_ids]

    # Получение ответов и неправильных ответов к вопросам
    answers = [file1.get_answer_from_db(question_id) for question_id in question_ids]
    wrong_answers = [file1.get_wrong_answer_from_db(question_id) for question_id in question_ids]

    # Получаем веса для каждого вопроса
    weights = [file1.get_question_weight_by_id(questid) for questid in question_ids]

    # Выводим весы в консоль для проверки (можете удалить эту строку после проверки)
    for questid, weight in zip(question_ids, weights):
        print(f"Вес вопроса с questid {questid}: {weight}")

    # Выводим id вопроса для каждого ответа в списке answers
    for answer, question_id in zip(answers, question_ids):
        print(f"Ответ: {answer}, Идентификатор вопроса: {question_id}")

    # Меняем порядок ответов и неправильных ответов для каждого вопроса
    for i in range(len(questions)):
        question_answers = answers[i] + wrong_answers[i]
        # random.shuffle(question_answers)
        answers[i], wrong_answers[i] = question_answers[:len(answers[i])], question_answers[len(answers[i]):]

    # Вывод вопросов и ответов на странице
    return render_template('test_form.html', theme=selected_theme, questions=questions, question_ids=question_ids,
                           answers=answers, wrong_answers=wrong_answers)


@app.route('/test_result')
def test_result():
    # Получаем questid из параметров запроса
    questid = request.args.get('questid')

    # # Получаем вес вопроса
    # weight = file1.get_question_weight_by_id(questid)
    #
    # # Выводим вес в консоль для проверки (можете удалить эту строку после проверки)
    # print(f"Вес вопроса с questid {questid}: {weight}")

    # Остальная логика вашей функции test_result
    return render_template('test_result.html', questid=questid, weight=weight)


# @app.route('/submit_answers', methods=['POST'])
# def submit_answers():
#     if request.method == 'POST':
#         # Инициализируем списки для ответов и неправильных ответов
#         answer_values = []
#         wrong_answer_values = []
#
#         for i in range(1, 6):  # 5 вопросов
#             answer_name = f'answer{i}'
#             wrong_answer_name = f'wrong_answer{i}'
#
#             # Получаем значения для текущего вопроса
#             answer_value = request.form.get(answer_name)
#             wrong_answer_value = request.form.get(wrong_answer_name)
#
#             # Добавляем значения в соответствующие списки, если они не являются None
#             if answer_value is not None:
#                 answer_values.append(answer_value)
#
#             if wrong_answer_value is not None:
#                 wrong_answer_values.append(wrong_answer_value)
#
#         # Получаем тему
#         selected_theme = request.args.get('theme')
#
#         # Получаем вопросы и их ID
#         questions_with_ids = file1.get_questions_from_db(selected_theme)
#         questions = [q[1] for q in questions_with_ids]
#         question_ids = [q[0] for q in questions_with_ids]
#
#
#         # Получаем значения из формы
#         selected_answer_count = len(request.form.getlist('answer'))
#         selected_wrong_answer_count = len(request.form.getlist('wrong_answer'))
#
#         # Возвращаем какой-то результат (можете редиректить на другую страницу)
#         return render_template('test_result.html',
#                                answer_values=answer_values,
#                                wrong_answer_values=wrong_answer_values,
#                                selected_answer_count=selected_answer_count,
#                                selected_wrong_answer_count=selected_wrong_answer_count,
#                                theme=selected_theme,
#                                questions=questions,
#                                question_ids=question_ids)


@app.route('/submit_answers', methods=['POST'])
def submit_answers():
    if request.method == 'POST':
        answer_values = []
        wrong_answer_values = []

        for i in range(1, 6):  # 5 вопросов
            answer_name = f'answer{i}'
            wrong_answer_name = f'wrong_answer{i}'

            process_answer(answer_name, answer_values)
            process_answer(wrong_answer_name, wrong_answer_values)

        selected_theme = session.get('theme')

        questions_with_ids = file1.get_questions_from_db(selected_theme)

        if questions_with_ids:
            questions, question_ids = zip(*[(q[1], q[0]) for q in questions_with_ids])
            print("questions_with_ids:")
            print(questions_with_ids)

        # else:
        #     # Handle the case where no questions are retrieved
        #     questions = []
        #     question_ids = []

        answers = [file1.get_answer_from_db(question_id) for question_id in question_ids]
        # print(answers)
        # print("wrong_answer_values")
        # print(wrong_answer_values)
        #
        wrong_answers = [file1.get_wrong_answer_from_db(question_id) for question_id in question_ids]
        wrong_answers_with_ids = [(question_id, file1.get_wrong_answer_from_db(question_id)) for question_id in
                                  question_ids]
        # print(wrong_answers_with_ids)

        wrong_answers_with_ids = list(zip(question_ids, wrong_answer_values))

        # Combine question text, correct answer, and wrong answer values with their corresponding question IDs
        combined_data = [(question_id, question_text, file1.get_answer_from_db(question_id), wrong_answer)
                         for (question_id, question_text), wrong_answer in zip(questions_with_ids, wrong_answer_values)]

        for question_id, question_text, correct_answer, wrong_answer in combined_data:
            print(f"ID вопроса: {question_id}")
            print(f"Текст вопроса: {question_text}")
            print(f"Правильный ответ: {correct_answer}")
            print(f"Неправильный ответ: {wrong_answer}")
            print("-" * 20)

        weights = [file1.get_question_weight_by_id(questid) for questid in question_ids]

        question_weights = dict(zip(question_ids, weights))

        matching_answers_with_ids = process_matching_answers()

        for answer_info in matching_answers_with_ids:
            answer_value, question_info = answer_info
            question_id, question_weight = question_info

            # Combine the information into a single string
            combined_output = f"Answer: {answer_value}, Question: {questions[question_ids.index(question_id)]}, Weight: {question_weight}"

            # print(combined_output)

        total_weight, mark = calculate_mark(matching_answers_with_ids)

        save_test_result(mark)

        # Вызов метода
        file1.calculate_average_scores()

        return render_template('test_result.html',
                               answer_values=answer_values,
                               wrong_answer_values=wrong_answer_values,

                               theme=selected_theme,
                               questions=questions,
                               question_ids=question_ids,
                               total_weight=total_weight,
                               mark=mark,
                               matching_answers_with_ids=matching_answers_with_ids,
                               answers=answers,
                               combined_output=combined_output,
                               combined_data=combined_data
                               )


def process_answer(answer_name, answer_values_list):
    answer_value = request.form.get(answer_name)
    if answer_value is not None:
        answer_values_list.append(answer_value)


def process_matching_answers():
    matching_answers_with_ids = []
    total_weight = 0

    for i in range(1, 6):  # 5 вопросов
        answer_name = f'answer{i}'
        answer_value = request.form.get(answer_name)

        if answer_value is not None:
            question_info = file1.get_question_id_by_answer(answer_value)

            matching_answers_with_ids.append((answer_value, question_info))
            total_weight += question_info[1]

    return matching_answers_with_ids


def calculate_mark(matching_answers_with_ids):
    total_weight = sum(weight for _, (_, weight) in matching_answers_with_ids)
    raw_mark = total_weight / 10
    mark = math.ceil(raw_mark)
    session['mark'] = mark
    return total_weight, mark


def save_test_result(mark):
    current_stud_id = session.get('current_stud_id')
    current_test_id = session.get('current_test_id')

    if current_stud_id is not None and current_test_id is not None:
        file1.insert_test_result(current_test_id, current_stud_id, mark)


# @app.route('/submit_answers', methods=['POST'])
# def submit_answers():
#     if request.method == 'POST':
#         # Инициализируем списки для ответов и неправильных ответов
#         answer_values = []
#         wrong_answer_values = []
#
#
#         for i in range(1, 6):  # 5 вопросов
#             answer_name = f'answer{i}'
#             wrong_answer_name = f'wrong_answer{i}'
#
#             # Получаем значения для текущего вопроса
#             answer_value = request.form.get(answer_name)
#             wrong_answer_value = request.form.get(wrong_answer_name)
#
#             # Добавляем значения в соответствующие списки, если они не являются None
#             if answer_value is not None:
#                 answer_values.append(answer_value)
#
#             if wrong_answer_value is not None:
#                 wrong_answer_values.append(wrong_answer_value)
#
#         # Получаем тему
#         selected_theme = request.args.get('theme')
#
#         # Получаем вопросы и их ID
#         questions_with_ids = file1.get_questions_from_db(selected_theme)
#         questions = [q[1] for q in questions_with_ids]
#         question_ids = [q[0] for q in questions_with_ids]
#
#         answers = [file1.get_answer_from_db(question_id) for question_id in question_ids]
#         wrong_answers = [file1.get_wrong_answer_from_db(question_id) for question_id in question_ids]
#
#         # Получаем веса для каждого вопроса
#         weights = [file1.get_question_weight_by_id(questid) for questid in question_ids]
#
#
#         # Создаем словарь для соотнесения вопросов с их весами
#         question_weights = dict(zip(question_ids, weights))
#
#         ######
#
#         matching_answers = []  # List to store matching answers
#         matching_answers_with_ids = []  # Список для хранения кортежей (answer_value, question_id)
#
#         # Получение ответов и неправильных ответов к вопросам
#         answers = [file1.get_answer_from_db(question_id) for question_id in question_ids]
#         wrong_answers = [file1.get_wrong_answer_from_db(question_id) for question_id in question_ids]
#
#
#         matching_answers_with_ids = []  # Список для хранения кортежей (answer_value, question_info)
#         total_weight = 0  # Переменная для хранения суммы весо
#         mark = 0
#
#         for i in range(1, 6):  # 5 вопросов
#             answer_name = f'answer{i}'
#
#             # Получаем значение для текущего вопроса
#             answer_value = request.form.get(answer_name)
#
#             # Добавляем значение в список, если оно не является None
#             if answer_value is not None:
#                 # Используем обновленный метод для получения информации о вопросе по значению ответа
#                 question_info = file1.get_question_id_by_answer(answer_value)
#
#                 # Добавляем (answer_value, question_info) в matching_answers_with_ids
#                 matching_answers_with_ids.append((answer_value, question_info))
#
#                 # Накапливаем вес в общую сумму
#                 total_weight += question_info[1]  # question_info[1] содержит вес
#
#                 # Пересчитываем оценку
#         raw_mark = total_weight / 10  # Получаем дробную часть от деления на 10
#         mark = math.ceil(raw_mark)  # Округляем вверх
#         # Сохраняем оценку и тему в сессии
#         session['mark'] = mark
#
#         # Выводим все непустые ответы
#         print("Matching Answers:", [answer_value for answer_value, _ in matching_answers_with_ids])
#
#         # Выводим все непустые ответы с соответствующими question_id и weight
#         for answer_value, (question_id, weight) in matching_answers_with_ids:
#             print(f"Совпавший ответ: {answer_value}, Идентификатор вопроса: {question_id}, Вес вопроса: {weight}")
#
#         # Выводим сумму весов всех ответов
#         print("Общий вес всех ответов:", total_weight)
#
#         current_stud_id = session.get('current_stud_id')
#         current_test_id = session.get('current_test_id')
#
#
#         # Проверка, что оба идентификатора доступны
#         if current_stud_id is not None and current_test_id is not None:
#             file1.insert_test_result(current_test_id, current_stud_id, mark)
#
#
#
#         # Получаем значения из формы
#         selected_answer_count = len(request.form.getlist('answer'))
#         selected_wrong_answer_count = len(request.form.getlist('wrong_answer'))
#
#         # Возвращаем какой-то результат (можете редиректить на другую страницу)
#         return render_template('test_result.html',
#                                answer_values=answer_values,
#                                wrong_answer_values=wrong_answer_values,
#                                selected_answer_count=selected_answer_count,
#                                selected_wrong_answer_count=selected_wrong_answer_count,
#                                theme=selected_theme,
#                                # answered_questions=answered_questions,
#                                # answered_weights=answered_weights,
#                                questions=questions,
#                                question_ids=question_ids,
#                                total_weight=total_weight,
#                                mark=mark
#         )


@app.route('/view_results', methods=['GET', 'POST'])
def view_results():
    current_stud_id = session.get('current_stud_id')

    if current_stud_id is not None:
        # Получение списка тем из базы данных
        themes = file1.get_themes_from_db()

        if request.method == 'POST':
            # Если получен POST-запрос, значит, пользователь выполняет поиск
            theme = request.form.get('theme')
            date = request.form.get('date')

            # Вызываем метод search_results_in_db для выполнения поиска
            search_test_results = file1.search_results_in_db(theme, date)
            print(search_test_results)

            # Передаем результаты поиска, тему, дату и список тем в шаблон
            return render_template('view_results.html', search_test_results=search_test_results, theme=theme, date=date, themes=themes)
        else:
            # Если запрос GET, то просто получаем результаты тестов
            test_results = file1.get_test_results_by_studid(current_stud_id)

            # Передаем результаты тестов и список тем в шаблон
            return render_template('view_results.html', test_results=test_results, themes=themes)
    else:
        return "Ошибка: Не удалось получить идентификатор студента из сессии."


@app.route('/save_to_word')
def download_word_report():
    current_stud_id = session.get('current_stud_id')

    if current_stud_id is not None:
        test_results = file1.get_test_results_by_studid(current_stud_id)
        if test_results:
            file_path = f"test_results_{current_stud_id}.docx"
            save_to_word(test_results, file_path)
            return send_file(file_path, as_attachment=True)
        else:
            return "Нет результатов для создания отчёта."
    else:
        return "Ошибка: Не удалось получить идентификатор студента из сессии."


@app.route('/search_save_to_word', methods=['GET'])
def search_save_to_word():
    current_stud_id = session.get('current_stud_id')

    if current_stud_id is not None:
        theme = request.args.get('theme')
        date = request.args.get('date')
        print(f"Search Parameters: Theme={theme}, Date={date}")

        search_test_results = file1.search_results_in_db(theme, date)

        if search_test_results:
            file_path = f"search_test_results_{current_stud_id}.docx"
            save_to_word(search_test_results, file_path)
            return send_file(file_path, as_attachment=True)
        else:
            return "Нет результатов для создания отчёта."
    else:
        return "Ошибка: Не удалось получить идентификатор студента из сессии."


def save_to_word(test_results, file_path='full_test_results.docx'):

    doc = Document()

    for result in test_results:
        doc.add_heading(result[5], level=2)
        doc.add_paragraph(f"Дата: {result[2]}")
        doc.add_paragraph(f"Результат: {result[3]}")
        doc.add_paragraph(f"Средняя оценка: {result[4]}")

        doc.add_paragraph('---')

    doc.save(file_path)

def search_save_to_word(search_test_results, file_path='search_test_results.docx'):

    doc = Document()
    print(f"Search Parameters2")

    # Добавляем данные из search_test_results в документ
    for result in search_test_results:
        doc.add_heading(result[5], level=2)
        doc.add_paragraph(f"Дата: {result[2]}")
        doc.add_paragraph(f"Результат: {result[3]}")
        doc.add_paragraph(f"Средняя оценка: {result[4]}")

        doc.add_paragraph('---')


    doc.save(file_path)


@app.route('/admin')
def admin():
    try:
        students_info = file1.get_all_students()
        themes, average_scores = calculate_average_scores()

        if students_info is not None:
            # Возвращаем шаблон admin.html с информацией о студентах и средних оценках
            return render_template('admin.html', students_info=students_info, themes=themes,
                                   average_scores=average_scores)
        else:
            # Обработка случая, если возникла ошибка при получении данных о студентах
            return render_template('error.html', error_message="Ошибка при получении данных о студентах")

    except Exception as e:
        # Обработка общих ошибок, которые могут возникнуть при выполнении кода
        return render_template('error.html', error_message=str(e))


@app.route('/students')
def display_students():
    students_info = file1.get_all_students()  # Получаем данные о студентах из вашей функции

    print("Information about students:")
    for student_info in students_info:
        print(student_info)

    if students_info:
        return render_template('admin.html', students_info=students_info)
    else:
        return "Произошла ошибка при получении данных о студентах."


def calculate_average_scores():
    try:
        # Получаем темы из базы данных
        themes = file1.get_themes_from_db()

        # Добавим расчет средних оценок для каждой темы
        average_scores = {theme: file1.average_score_by_theme(theme) for theme in themes}
        # print(themes, average_scores)
        return themes, average_scores

    except Exception as e:
        print(f"Ошибка при расчете средних оценок: {e}")
        return None, None

def test_count_by_theme_plot():
    # Получение данных о количестве тестов по темам
    test_count_by_theme = file1.get_test_count_by_theme()

    # Получение списка тем
    themes = file1.get_themes_from_db()

    # Создаем словарь для связи тем и количества тестов
    theme_count_dict = dict(test_count_by_theme)

    # Получаем количество тестов для каждой темы
    counts = [theme_count_dict.get(theme, 0) for theme in themes]

    # Создание графика
    plt.bar(themes, counts)
    plt.xlabel('Темы')
    plt.ylabel('Количество тестов')
    plt.title('Количество тестов по темам')
    plt.xticks(rotation=20, ha='right')
    plt.tight_layout()

    # Сохраняем график в байтовый объект
    image_stream = BytesIO()
    plt.savefig(image_stream, format='png')

    # Ensure image_stream is seekable
    image_stream.seek(0)

    # Кодируем изображение в base64
    image_base64 = base64.b64encode(image_stream.read()).decode('utf-8')

    plt.clf()
    plt.close()

    return image_base64

@app.route('/theme_scores_plot')
def theme_scores_plot():
    # Получение данных о темах
    themes = file1.get_themes_from_db()

    # Создание словаря средних оценок
    average_scores = {theme: file1.average_mark_by_theme(theme) for theme in themes}

    # print(average_scores)
    # Фильтрация None-значений (если есть)
    average_scores = {theme: score for theme, score in average_scores.items() if score is not None}

    # Создание графика
    plt.bar(average_scores.keys(), list(average_scores.values()))
    plt.xlabel('Темы')
    plt.ylabel('Средняя оценка')
    plt.title('Средние оценки по темам')
    plt.xticks(rotation=20, ha='right')
    # Установка масштаба на оси Y от 0 до 10
    plt.ylim(0, 10)
    plt.tight_layout()

    # Сохраняем график в байтовый объект
    image_stream = BytesIO()
    plt.savefig(image_stream, format='png')
    image_stream.seek(0)

    # Кодируем изображение в base64
    image_base64 = base64.b64encode(image_stream.read()).decode('utf-8')

    # Очищаем текущий график
    plt.clf()
    second_image_base64 = test_count_by_theme_plot()

    pie_image_base64 = students_per_theme_plot()

    # Отображаем шаблон с изображением
    return render_template(
        'theme_scores_plot.html',
        image_base64=image_base64,
        second_image_base64=second_image_base64,
        pie_image_base64=pie_image_base64
    )



@app.route('/students_per_theme_plot')
def students_per_theme_plot():
    # Получение данных о количестве студентов по темам
    students_per_theme = file1.get_students_count_by_theme()

    # Разделение данных на темы и количество студентов
    themes = [theme for theme, count in students_per_theme]
    counts = [count for theme, count in students_per_theme]

    # Создание круговой диаграммы
    plt.figure(figsize=(8, 8))
    plt.pie(counts, labels=themes, autopct='%1.1f%%', startangle=90, colors=plt.cm.Paired.colors)
    plt.title('Процентное распределение студентов, прошедших тесты по темам')
    plt.axis('equal')  # Делает круговой график кругом

    # Сохранение графика в байтовый объект
    image_stream = BytesIO()
    plt.savefig(image_stream, format='png')
    image_stream.seek(0)

    # Кодируем изображение в base64
    pie_image_base64 = base64.b64encode(image_stream.read()).decode('utf-8')

    # Очищаем текущий график
    plt.clf()
    plt.close()

    # Возвращаем base64-картинку
    return pie_image_base64




# @app.route('/pie_chart')
# def pie_chart():
#     try:
#         # Получите данные о темах и количестве тестов
#         test_count_by_theme = file1.get_test_count_by_theme()
#
#         # Печать результата
#         print("Результат запроса по количеству тестов по темам:", test_count_by_theme)
#
#         themes = file1.get_themes_from_db()
#         test_count_by_theme = file1.get_test_count_by_theme()
#
#         # Создаем словарь, чтобы связать темы и количество тестов
#         theme_count_dict = dict(test_count_by_theme)
#
#         # Получаем количество тестов для каждой темы
#         counts = [theme_count_dict.get(theme, 0) for theme in themes]
#         for theme, count in zip(themes, counts):
#             print(f"Темааа '{theme}': {count} тестов")
#
#         # Строим график
#         plt.figure(figsize=(10, 6))
#         plt.bar(themes, counts, color='skyblue')
#         plt.xlabel('Темы')
#         plt.ylabel('Количество тестов')
#         plt.title('Количество тестов по темам')
#         plt.xticks(rotation=45, ha='right')
#
#
#         # Выводим график
#         plt.tight_layout()
#         plt.show()
#
#         # Сохраняем график в байтовый объект
#         image_stream = BytesIO()
#         plt.savefig(image_stream, format='png')
#         image_stream.seek(0)
#
#         # Кодируем изображение в base64
#         image_base64_pie = base64.b64encode(image_stream.read()).decode('utf-8')
#
#
#
#         # Возвращаем HTML-код для отображения в шаблоне
#         return render_template('theme_scores_plot.html', image_base64_pie=image_base64_pie)
#
#     except Exception as e:
#         print(f"Ошибка при построении круговой диаграммы: {e}")
#         return "Ошибка при построении круговой диаграммы"


# @app.route('/pie_chart')
# def pie_chart():
#     try:
#         # Получите данные о темах и количестве тестов
#         test_count_by_theme = file1.get_test_count_by_theme()
#
#         # Печать результата
#         print("Результат запроса по количеству тестов по темам:", test_count_by_theme)
#
#         themes = file1.get_themes_from_db()
#         test_count_by_theme = file1.get_test_count_by_theme()
#
#         # Создаем словарь, чтобы связать темы и количество тестов
#         theme_count_dict = dict(test_count_by_theme)
#
#         # Получаем количество тестов для каждой темы
#         counts = [theme_count_dict.get(theme, 0) for theme in themes]
#         for theme, count in zip(themes, counts):
#             print(f"Темааа '{theme}': {count} тестов")
#         plt.show()
#         # Создание круговой диаграммы
#         plt.figure(figsize=(8, 8))
#         plt.pie(counts, labels=themes, autopct='%1.1f%%', startangle=140)
#         plt.title('Количество тестов по темам')
#
#         # Сохраняем график в байтовый объект
#         image_stream = BytesIO()
#         plt.savefig(image_stream, format='png')
#         image_stream.seek(0)
#
#         # Кодируем изображение в base64
#         image_base64_pie = base64.b64encode(image_stream.read()).decode('utf-8')
#
#
#
#         # Возвращаем HTML-код для отображения в шаблоне
#         return render_template('theme_scores_plot.html', image_base64_pie=image_base64_pie)
#
#     except Exception as e:
#         print(f"Ошибка при построении круговой диаграммы: {e}")
#         return "Ошибка при построении круговой диаграммы"



from flask import render_template

@app.route('/add_student', methods=['POST'])
def add_student():
    # Получаем данные из формы
    firstname = request.form.get('firstname')
    lastname = request.form.get('lastname')
    course = request.form.get('course')
    groupnumber = request.form.get('groupnumber')

    # Проверяем, существует ли студент
    if file1.check_student_in_db(firstname, lastname):
        return render_template('admin.html', student_exists=True)
    else:
        # Если студент не существует, добавляем его в базу данных
        if file1.add_student_to_db(firstname, lastname, course, groupnumber):
            return redirect(url_for('admin'))  # Перенаправляем на страницу администратора после успешного добавления
        else:
            return "Ошибка при добавлении студента в базу данных"


@app.route('/delete_student', methods=['POST'])
def delete_student():
    # Получите id студента из запроса
    student_id = request.form.get('student_id')
    print(student_id)
    # Удалите студента из базы данных с использованием новой функции
    if file1.delete_student_by_id(student_id):
        return redirect(url_for('admin'))  # Перенаправляем на страницу администратора после успешного удаления
    else:
        return "Ошибка при удалении студента"



@app.route('/add_question', methods=['POST'])
def add_question():
    if request.method == 'POST':
        # Получаем данные из формы
        question_text = request.form.get('question_text')
        answer = request.form.get('answer')
        weight = request.form.get('weight')
        theme = request.form.get('theme')
        wrong_answer = request.form.get('wrong_answer')

        # Проверяем существование вопроса в базе данных
        question_exists = file1.check_question_in_db(question_text)

        # Если вопрос существует, передаем информацию в шаблон
        if question_exists:
            return render_template('admin.html', question_exists=True)

        # Добавляем студента в базу данных
        if file1.add_question_to_db(question_text, answer, weight, theme, wrong_answer):
            return redirect(url_for('admin'))  # Перенаправляем на страницу администратора после успешного добавления
        else:
            return "Ошибка при добавлении вопроса в базу данных"


@app.route('/download_excel', methods=['GET'])
def download_excel():
    # Получите данные о студентах (students_info)
    students_data = file1.get_all_students()

    # Создайте новую книгу Excel
    wb = openpyxl.Workbook()
    ws = wb.active

    headers = ['student_id', 'firstname', 'lastname', 'course', 'groupnumber']
    headers_mapping = {
        'student_id': 'ID',
        'firstname': 'Имя',
        'lastname': 'Фамилия',
        'course': 'Курс',
        'groupnumber': 'Номер группы'
    }

    # Добавляем данные студентов
    for row_num, student in enumerate(students_data, 2):
        for col_num, key in enumerate(headers, 1):
            value = student.get(key, '')
            ws.cell(row=row_num, column=col_num, value=value)

    # Заменяем заголовки в первой строке
    for col_num, key in enumerate(headers, 1):
        ws.cell(row=1, column=col_num, value=headers_mapping.get(key, key))

    # # Заголовки
    # headers = ['student_id', 'firstname', 'lastname', 'course', 'groupnumber']
    #
    # for col_num, header in enumerate(headers, 1):
    #     ws.cell(row=1, column=col_num, value=header)
    #
    # # Добавляем данные студентов
    # for row_num, student in enumerate(students_data, 2):
    #     print(f"Processing student: {student}")
    #
    #     for col_num, key in enumerate(headers, 1):
    #         value = student.get(key, '')
    #         print(f"Setting key {key} with value {value} in cell ({row_num}, {col_num})")
    #
    #         ws.cell(row=row_num, column=col_num, value=value)

    # Сохраняем книгу Excel
    excel_file_path = 'students_data.xlsx'
    wb.save(excel_file_path)

    return send_from_directory('.', excel_file_path, as_attachment=True)





@app.route('/export', methods=['POST'])
def export_to_json_route():
    conn = psycopg2.connect(user="postgres", password="rootparol", host="127.0.0.1", port="5432", database="MyDB")
    cursor = conn.cursor()

    filename = 'database_file.json'

    # Ваш SQL-запрос для объединения данных из всех таблиц
    sql_queries = [
        "SELECT * FROM users",
        "SELECT * FROM students",
        "SELECT * FROM professors",
        "SELECT * FROM tests",
        "SELECT * FROM questions",
    ]

    data_from_db = []

    for sql_query in sql_queries:
        cursor.execute(sql_query)
        columns = [column[0] for column in cursor.description]
        data_from_db.extend([dict(zip(columns, row)) for row in cursor.fetchall()])

    # Закрываем соединение
    conn.close()

    # Создаем JSON файл с данными
    with open(filename, 'w') as file:
        json.dump(data_from_db, file)

    # Отправляем файл для скачивания
    return send_file(filename, as_attachment=True)


import psycopg2


@app.route('/import', methods=['POST'])
def insert_data():
    table_name = request.form.get(
        'table_name')  # Получаем имя таблицы из запроса (предположим, что данные передаются через форму)
    data = request.form.get(
        'data')  # Получаем данные из запроса (предположим, что данные передаются через форму и в формате JSON)

    if not table_name or not data:
        return "Не указано имя таблицы или данные."

    # Преобразуем строку JSON в объект Python
    data = json.loads(data)

    conn = psycopg2.connect(user="postgres", password="rootparol", host="127.0.0.1", port="5432", database="MyDB")
    cursor = conn.cursor()

    try:
        # Получаем информацию о столбцах таблицы
        cursor.execute(
            f"SELECT COLUMN_NAME FROM INFORMATION_SCHEMA.COLUMNS WHERE TABLE_NAME = '{table_name}' AND COLUMNPROPERTY(object_id(TABLE_NAME), COLUMN_NAME, 'IsIdentity') != 1")
        non_identity_columns = [column[0] for column in cursor.fetchall()]

        # Формируем строку SQL-запроса для вставки данных
        columns_without_identity = ', '.join(non_identity_columns)
        values_placeholder = ', '.join(['%s' for _ in range(len(non_identity_columns))])
        sql_query = f"INSERT INTO {table_name} ({columns_without_identity}) VALUES ({values_placeholder})"

        # Вставляем данные в таблицу
        for row in data:
            # Передаем только значения, соответствующие не-identity столбцам
            values = tuple(row[column] for column in non_identity_columns)

            # Выполняем запрос
            cursor.execute(sql_query, values)

        conn.commit()
        return "Данные успешно импортированы в базу данных."
    except Exception as e:
        conn.rollback()
        return f"Ошибка импорта данных: {str(e)}"
    finally:
        conn.close()

#####

def insert_data(table_name, data):
    conn = psycopg2.connect(user="postgres", password="rootparol", host="127.0.0.1", port="5432", database="MyDB")
    cursor = conn.cursor()

    # Получаем информацию о столбцах таблицы
    cursor.execute(
        f"SELECT COLUMN_NAME FROM INFORMATION_SCHEMA.COLUMNS WHERE TABLE_NAME = '{table_name}' AND COLUMN_NAME NOT IN (SELECT column_name FROM information_schema.columns WHERE table_name = '{table_name}' AND column_default LIKE 'nextval%')"
    )
    non_identity_columns = [column[0] for column in cursor.fetchall()]

    # Формируем строку SQL-запроса для вставки данных
    columns_without_identity = ', '.join(non_identity_columns)
    values_placeholder = ', '.join(['?' for _ in range(len(non_identity_columns))])
    sql_query = f"INSERT INTO {table_name} ({columns_without_identity}) VALUES ({values_placeholder})"

    try:
        # Вставляем данные в таблицу
        for row in data:
            # Передаем только значения, соответствующие не-identity столбцам
            values = [row[column] for column in non_identity_columns]

            # Выполняем запрос
            cursor.execute(sql_query, values)

        conn.commit()
        return "Данные успешно импортированы в базу данных."
    except Exception as e:
        conn.rollback()
        return f"Ошибка импорта данных: {str(e)}"
    finally:
        conn.close()




@app.route('/import_data', methods=['POST'])
def import_data():
    # Получаем файл из формы
    uploaded_file = request.files['file']
    table_name = request.form.get('table_name')
    data = json.loads(request.form.get('data'))

    # Чтение данных из JSON-файла
    json_data = json.load(uploaded_file)

    # Импорт данных в каждую таблицу
    insert_data('students', json_data)
    # insert_data('Airports', json_data)
    # insert_data('Planes', json_data)
    # insert_data('Flights', json_data)
    # insert_data('Users', json_data)
    # insert_data('Passengers', json_data)
    # insert_data('Tickets', json_data)
    # insert_data('TicketPurchaseHistory', json_data)

    return "Данные успешно импортированы в базу данных."


@app.route('/importjson')
def importjson():
    return render_template('importjson.html')









# @app.route('/all_students')
# def all_students():
#     students_info = file1.get_all_students()
#
#     if students_info:
#         return render_template('all_students.html', students_info=students_info)
#     else:
#         return "Нет доступной информации о студентах"


if __name__ == "__main__":
    app.run(debug=True)
#
# #
# #
# # # bcrypt.check_password_hash(user_data[2], password)
# #
#
# # from flask import session
# # import file1
#
# # # Получаем username из сессии
# # username = session.get('username')
# # print(f"""Вошли с данными {username}""")
# # # Получаем информацию о студенте
# # student_info_dict = file1.get_student_info('username')
# #
# # # Вывести содержимое словаря в консоль
# # print("Информация о студенте:")
# # for key, value in student_info_dict.items():
# #     print(f"{key}: {value}")
#

import file1

# Получаем список тем
# themes = file1.get_themes_from_db()
# for theme in themes:
#     print(theme)

# # Получаем вопросы для выбранной темы
# selected_theme = 'Языкознание и литература'
# questions = file1.get_questions_from_db(selected_theme)
#
# # Выводим информацию о вопросах
# for question in questions:
#     print(question)
#
# # Выводим информацию о вопросах и их ответах
# for question in questions:
#     print(f"Вопрос: {question}")
#     quest_id = question['quest_id']
#     answer = file1.get_answer_from_db(quest_id)
#     print(f"Ответ: {answer}")
#     print("---")

# students_info = file1.get_all_students()
#
#
# for student in students_info:
#     print(f" Имя:{student['firstname']}, Фамилия: {student['lastname']}, Курс: {student['course']}, Группа: {student['groupnumber']}")
#
#     # 'firstname': student[1],
#     # 'lastname': student[2],
#     # 'course': student[3],
#     # 'groupnumber': student[4],
